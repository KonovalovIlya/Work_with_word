from docx import Document
import os


class Data:

    def __init__(self, number, date, side, company, director):
        self.number = number
        self.date = date
        self.side = side
        self.company = company
        self.director = director

    def get_number(self):
        return self.number

    def get_date(self):
        return self.date

    def get_side(self):
        return self.side

    def get_company(self):
        return self.company

    def get_director(self):
        return self.director


dir = os.path.abspath('.')
for j, i in enumerate(os.listdir(dir)):
    if i.startswith('.') or i.startswith('Дополнительное') or i == 'file.py':
        pass
    else:
        # print(i)
        doc = Document(i)
        number = doc.paragraphs[0].runs[4].text

        # date =


def make_data(file_name):
    doc = Document(file_name)
    doc.paragraphs[2].runs

