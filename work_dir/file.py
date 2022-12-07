from docx import Document
import os

from docx.shared import Pt


class Data:

    def __init__(self, number, date, side, company, director):
        self.number = number
        self.date = date
        self.side = side
        self.company = company
        self.director = director

    def get_number(self):
        return self.number

    def get_date(self):
        return self.date

    def get_side(self):
        return self.side

    def get_company(self):
        return self.company

    def get_director(self):
        return self.director


def make_dops(data_dict):
    dir = os.path.abspath('.')
    for i in os.listdir(dir):
        if i.startswith('draft'):
            doc = Document(i)
            doc.paragraphs[1].text = ''.join(
                [doc.paragraphs[1].runs[0].text, data_dict.get('number')[0]]
            )
            for run in doc.paragraphs[1].runs:
                run.font.name = 'Calibri'
                run.font.bold = True
            doc.paragraphs[2].text = data_dict.get('date_129')[0]
            for run in doc.paragraphs[2].runs:
                run.font.name = 'Calibri'
                run.font.bold = True
            doc.paragraphs[6].text = ''
            for run in data_dict.get('side')[0]:
                doc.paragraphs[6].add_run(run.text)
            doc.paragraphs[6].runs[1].font.size = Pt(10)
            doc.paragraphs[6].runs[1].font.name = 'Calibri'
            doc.paragraphs[6].runs[1].font.bold = True
            for run in doc.paragraphs[6].runs[2:]:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
            runs = doc.paragraphs[8].text.split(' ')
            doc.paragraphs[8].text = ''
            for ii, run in enumerate(runs):
                doc.paragraphs[8].add_run(f'{run} ')
                doc.paragraphs[8].runs[ii+1].font.name = 'Calibri'
                doc.paragraphs[8].runs[ii+1].font.size = Pt(10)
            print(doc.paragraphs[1].runs[0].text, '\n', doc.paragraphs[2].text)
            to_save = f'res/Dopnik k dogovoru №{data_dict["number"][0]}.docx'
            doc.save(to_save)


def make_data():
    dir = os.path.abspath('.')
    data_dict = {
        'number': [],
        'date_129': [],
        'side': [],
        'director': []
    }
    for j, i in enumerate(os.listdir(dir)):
        if not os.path.isdir(i) and not i.startswith('Дополнительное') and not i.startswith('draft') and i != 'file.py':
            doc = Document(i)
            data_dict['number'].append(doc.paragraphs[0].runs[4].text)
            data_dict['date_129'].append(doc.paragraphs[129].text)
            data_dict['side'].append(doc.paragraphs[2].runs)
            data_dict['director'].append(doc.paragraphs[2].runs[8:13])
    make_dops(data_dict)


# if __name__ == 'main':
make_data()
